"""
document_parser.py

Milestone 3A — JD + Resume ingestion.

Turns a raw input document (.txt, .md, .pdf, .docx) into a normalized,
plain-text internal representation the rest of the pipeline can use later.

This module does NOT call any LLM and does NOT talk to Gemini/LiveKit.
It is intentionally isolated from src/realtime/ — the realtime agent
does not import from here yet.
"""

from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class ParsedDocument:
    doc_type: str          # "job_description" | "resume" | "other"
    filename: str          # original filename
    source_format: str     # ".txt" | ".md" | ".pdf" | ".docx"
    text: str               # normalized plain text
    char_count: int

    def to_dict(self) -> dict:
        return asdict(self)


def _normalize_text(raw_text: str) -> str:
    """
    Collapses excessive whitespace/blank lines so downstream parsing
    (later: LLM structured extraction) sees clean, consistent text
    regardless of source format.
    """
    lines = [line.strip() for line in raw_text.splitlines()]
    lines = [line for line in lines if line != ""]
    return "\n".join(lines).strip()


def _read_txt_or_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def parse_document(file_path: str | Path, doc_type: str = "other") -> ParsedDocument:
    """
    Parses a single input document into a ParsedDocument.

    Args:
        file_path: path to a .txt, .md, .pdf, or .docx file
        doc_type: "job_description" or "resume" (free label, not validated
                   here — kept simple for this milestone)
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Input document not found: {path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if ext in (".txt", ".md"):
        raw_text = _read_txt_or_md(path)
    elif ext == ".pdf":
        raw_text = _read_pdf(path)
    else:  # ".docx"
        raw_text = _read_docx(path)

    normalized = _normalize_text(raw_text)

    return ParsedDocument(
        doc_type=doc_type,
        filename=path.name,
        source_format=ext,
        text=normalized,
        char_count=len(normalized),
    )